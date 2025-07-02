import os
import re
from telegram import Update
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes
from docx import Document

# --- Paste your bot token here ---
TELEGRAM_BOT_TOKEN = '8077706019:AAE5qdQ6i4IyNkTAzxCiBv-45xFJSnCWD9o'

def create_docx(questions_data, file_path):
    """
    Generates a .docx file with a separate, fixed-structure table for each question.
    This version includes more robust checking for the correct option.
    """
    doc = Document()
    for q_data in questions_data:
        table = doc.add_table(rows=0, cols=3)
        table.style = 'Table Grid'
        
        # --- 1. Question Row ---
        row_cells = table.add_row().cells
        row_cells[0].text = 'Question'
        row_cells[1].merge(row_cells[2]).text = q_data['question_text']

        # --- 2. Type Row ---
        row_cells = table.add_row().cells
        row_cells[0].text = 'Type'
        row_cells[1].merge(row_cells[2]).text = 'multiple_choice'

        # --- 3. Option Rows ---
        correct_index_val = q_data.get('correct_option_index')
        
        # Defensive check to ensure the index is a valid integer
        try:
            correct_index = int(correct_index_val)
        except (ValueError, TypeError):
            correct_index = -1 # Set to an impossible index if invalid

        for i, option in enumerate(q_data['options']):
            row_cells = table.add_row().cells
            row_cells[0].text = 'Option'
            row_cells[1].text = option['text'] # Options are now dicts
            row_cells[2].text = 'correct' if i == correct_index else 'incorrect'

        # --- 4. Solution Row ---
        row_cells = table.add_row().cells
        row_cells[0].text = 'Solution'
        row_cells[1].merge(row_cells[2]).text = q_data.get('explanation_text', '')

        # --- 5. Marks Row ---
        row_cells = table.add_row().cells
        row_cells[0].text = 'Marks'
        row_cells[1].text = '4'
        row_cells[2].text = '1'
        doc.add_paragraph('')
    doc.save(file_path)

async def process_and_send_docx(chat_id: int, context: ContextTypes.DEFAULT_TYPE):
    """Helper function to generate and send the DOCX file, then clear the queue."""
    questions_queue = context.user_data.get('questions_queue', [])
    if not questions_queue:
        await context.bot.send_message(chat_id, "There are no questions in the queue to convert.")
        return

    await context.bot.send_message(chat_id, f"Processing {len(questions_queue)} question(s)... \U0001F4DD")
    file_path = f'questions_{chat_id}.docx'
    
    try:
        create_docx(questions_queue, file_path)
        await context.bot.send_message(chat_id, f"\u2705 Successfully created the .docx file.")
        await context.bot.send_document(chat_id, document=open(file_path, 'rb'))
        os.remove(file_path)
    except Exception as e:
        await context.bot.send_message(chat_id, f"\U0001F614 Sorry, an error occurred while creating the .docx file: {e}")
    finally:
        context.user_data['questions_queue'] = []

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler for the /start command."""
    context.user_data.clear()
    await update.message.reply_text(
        "Hello! \U0001F44B\n\n"
        "I'm ready to collect your quizzes.\n\n"
        "**Workflow:**\n"
        "1. Forward a quiz.\n"
        "2. Forward the solution text for that quiz.\n\n"
        "I will automatically create a file after 30 quizzes, or you can use /convert at any time."
    )

async def handle_quiz(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handles incoming quizzes, storing them temporarily in a structured way."""
    poll = update.message.poll
    if poll.type != 'quiz':
        await update.message.reply_text("That was a poll, not a quiz. I can only process quizzes.")
        return

    if context.user_data.get('last_quiz'):
        pending_quiz = context.user_data['last_quiz']
        pending_quiz['explanation_text'] = ""
        context.user_data.setdefault('questions_queue', []).append(pending_quiz)
        queue_len = len(context.user_data['questions_queue'])
        await update.message.reply_text(f"⚠️ A previous quiz was missing a solution. It has been saved with a blank solution. Current queue: {queue_len}/30")

    # Store quiz data with options as dicts for consistency
    context.user_data['last_quiz'] = {
        'question_text': poll.question,
        'options': [{'text': opt.text} for opt in poll.options],
        'correct_option_index': poll.correct_option_id,
    }
    await update.message.reply_text("Got the quiz! Now waiting for the solution text... \U0001F447")

async def handle_solution_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handles the text message following a quiz, pairs it, and adds to the queue."""
    last_quiz = context.user_data.get('last_quiz')
    if not last_quiz:
        await update.message.reply_text(f"ERROR \U0001F6AB\n\nI received a text message, but I was expecting a quiz first. Please send a quiz before its solution.\n\n**Problematic Text:**\n{update.message.text}")
        return

    last_quiz['explanation_text'] = update.message.text
    questions_queue = context.user_data.setdefault('questions_queue', [])
    questions_queue.append(last_quiz)
    context.user_data['last_quiz'] = None
    
    queue_len = len(questions_queue)
    await update.message.reply_text(f"✅ Quiz and solution saved. ({queue_len}/30)")

    if queue_len >= 30:
        await process_and_send_docx(update.message.chat_id, context)

async def convert_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Manually triggers the conversion of the current queue."""
    if context.user_data.get('last_quiz'):
        pending_quiz = context.user_data['last_quiz']
        pending_quiz['explanation_text'] = ""
        context.user_data.setdefault('questions_queue', []).append(pending_quiz)
        context.user_data['last_quiz'] = None
        await update.message.reply_text("⚠️ A pending quiz was saved with a blank solution before converting.")

    await process_and_send_docx(update.message.chat_id, context)

async def cancel_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Clears the current question queue and any pending quiz."""
    context.user_data.clear()
    await update.message.reply_text("All pending quizzes and the current queue have been cleared. \U0001F5D1️")

def main():
    """Starts the bot and adds all handlers."""
    application = Application.builder().token(TELEGRAM_BOT_TOKEN).build()
    
    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("convert", convert_command))
    application.add_handler(CommandHandler("cancel", cancel_command))
    
    application.add_handler(MessageHandler(filters.POLL, handle_quiz))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_solution_text))
    
    print("Bot started... (Press Ctrl+C to stop)")
    application.run_polling()

if __name__ == '__main__':
    main()
